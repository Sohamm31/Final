# backend/app/services/interview_logic.py

import os
import re
import shutil
import tempfile
import json
import PyPDF2
import docx
from typing import List, Dict, Any, Tuple
from langchain_openai import ChatOpenAI
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import PydanticOutputParser
from pydantic import BaseModel, Field
from langchain.text_splitter import RecursiveCharacterTextSplitter, Language
from langchain_chroma import Chroma
from langchain_community.document_loaders import GitLoader
from langchain_core.documents import Document

# --- Configuration ---
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY", "")
llm = ChatOpenAI(
    model_name="deepseek/deepseek-r1-0528-qwen3-8b:free",
    openai_api_key=OPENROUTER_API_KEY,
    openai_api_base="https://openrouter.ai/api/v1",
    temperature=0.7,
    request_timeout=60
)
embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

# --- Pydantic Models for LLM Output ---
class TechnicalFeedback(BaseModel):
    technical_knowledge_rating: int = Field(description="Rating of technical knowledge out of 5.")
    technical_tips: List[str] = Field(description="List of actionable tips to improve technical answers.")

class HRFeedback(BaseModel):
    communication_skills_rating: int = Field(description="Rating of communication skills out of 5.")
    communication_tips: List[str] = Field(description="List of actionable tips to improve communication.")

class InterviewQuestion(BaseModel):
    question: str = Field(description="A single, direct interview question for the candidate.")

question_parser = PydanticOutputParser(pydantic_object=InterviewQuestion)

# --- Resume and Git Processing ---
def extract_text_from_pdf(file_path: str) -> str:
    text = ""
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            text += page.extract_text() or ""
    return text

def extract_text_from_docx(file_path: str) -> str:
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_github_links(text: str) -> List[str]:
    pattern = r'https?://github\.com/([a-zA-Z0-9_-]+/[a-zA-Z0-9_.-]+)'
    links = re.findall(pattern, text)
    return list(set([f"https://github.com/{link}" for link in links]))

async def clone_and_process_repo(repo_url: str) -> List[Document]:
    temp_dir = tempfile.mkdtemp()
    try:
        loader = GitLoader(repo_path=temp_dir, clone_url=repo_url, branch="main")
        docs = loader.load()
        
        python_splitter = RecursiveCharacterTextSplitter.from_language(language=Language.PYTHON, chunk_size=1000, chunk_overlap=200)
        # Add other splitters as needed
        
        processed_docs = []
        for doc in docs:
            if doc.metadata.get('file_path', '').endswith('.py'):
                processed_docs.extend(python_splitter.split_documents([doc]))
        return processed_docs
    finally:
        shutil.rmtree(temp_dir)

async def process_resume_and_embed(file_path: str, filename: str, session_id: str) -> Tuple[str, List[str], str]:
    file_ext = os.path.splitext(filename)[1].lower()
    text = ""
    if file_ext == '.pdf':
        text = extract_text_from_pdf(file_path)
    elif file_ext == '.docx':
        text = extract_text_from_docx(file_path)
    
    if not text:
        raise ValueError("Could not extract text from resume.")
        
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    resume_docs = [Document(page_content=chunk, metadata={"source": "resume"}) for chunk in text_splitter.split_text(text)]
    
    chroma_db_path = os.path.join(tempfile.gettempdir(), f"chroma_db_{session_id}")
    
    all_docs = resume_docs
    github_links = extract_github_links(text)
    github_knowledge = []
    for link in github_links:
        repo_docs = await clone_and_process_repo(link)
        all_docs.extend(repo_docs)
        github_knowledge.append(link.split('/')[-1])

    Chroma.from_documents(documents=all_docs, embedding=embeddings, persist_directory=chroma_db_path)
    
    return text, github_knowledge, chroma_db_path

# --- Interview Logic ---
async def _get_relevant_context(chroma_db_path: str, query: str) -> str:
    if not chroma_db_path: return ""
    vector_store = Chroma(persist_directory=chroma_db_path, embedding_function=embeddings)
    retrieved_docs = vector_store.similarity_search(query, k=3)
    return "\n---\n".join([doc.page_content for doc in retrieved_docs])

async def generate_question(session_data: Dict[str, Any], question_type: str, user_answer_prev: str = "") -> str:
    session_data['section_questions_asked'][question_type] += 1
    if session_data['section_questions_asked'][question_type] >= session_data['max_questions_per_section']:
        session_data['sections_covered'][question_type] = True

    context = await _get_relevant_context(session_data['chroma_db_path'], f"candidate's {question_type}")
    
    # *** FIX: Made the prompt much more strict to force JSON output ***
    prompt_template = """
    You are an AI assistant that ONLY responds with a valid JSON object. Do not add any other text, explanations, or markdown formatting.
    Your response must conform to the following JSON schema:
    {format_instructions}

    Now, based on the context below, generate one interview question.
    Question Type: {question_type}
    Resume Context: {context}
    Candidate's Previous Answer: {user_answer_prev}
    """
    
    prompt = PromptTemplate(
        template=prompt_template, 
        input_variables=["question_type", "context", "user_answer_prev"], 
        partial_variables={"format_instructions": question_parser.get_format_instructions()}
    )
    chain = prompt | llm | question_parser
    
    response = await chain.ainvoke({
        "question_type": question_type,
        "context": context,
        "user_answer_prev": user_answer_prev
    })
    
    session_data['conversation_history'].append(('ai', response.question))
    return response.question

def determine_next_question_type(session_data: Dict[str, Any]) -> str:
    if not session_data['sections_covered']['introduction']: return 'introduction'
    if not session_data['sections_covered']['skills']: return 'skills'
    if not session_data['sections_covered']['projects']: return 'projects'
    if not session_data['sections_covered']['experience']: return 'experience'
    return 'feedback_stage'

async def get_feedback(session_data: Dict[str, Any]) -> Tuple[TechnicalFeedback, HRFeedback]:
    full_conversation = "\n".join([f"{role.upper()}: {text}" for role, text in session_data['conversation_history']])
    
    # *** FIX: Made the feedback prompts more strict to force JSON output ***
    
    # Technical Feedback
    tech_parser = PydanticOutputParser(pydantic_object=TechnicalFeedback)
    tech_prompt_template = """
    You are an AI assistant that ONLY responds with a valid JSON object. Do not add any other text, explanations, or markdown formatting.
    Your response must conform to the following JSON schema:
    {format_instructions}

    Now, analyze the following conversation and provide technical feedback.
    Conversation:
    {conversation}
    """
    tech_prompt = PromptTemplate(
        template=tech_prompt_template, 
        input_variables=["conversation"], 
        partial_variables={"format_instructions": tech_parser.get_format_instructions()}
    )
    tech_chain = tech_prompt | llm | tech_parser
    technical_feedback = await tech_chain.ainvoke({"conversation": full_conversation})

    # HR Feedback
    hr_parser = PydanticOutputParser(pydantic_object=HRFeedback)
    hr_prompt_template = """
    You are an AI assistant that ONLY responds with a valid JSON object. Do not add any other text, explanations, or markdown formatting.
    Your response must conform to the following JSON schema:
    {format_instructions}

    Now, analyze the following conversation and provide HR/communication feedback.
    Conversation:
    {conversation}
    """
    hr_prompt = PromptTemplate(
        template=hr_prompt_template, 
        input_variables=["conversation"], 
        partial_variables={"format_instructions": hr_parser.get_format_instructions()}
    )
    hr_chain = hr_prompt | llm | hr_parser
    hr_feedback = await hr_chain.ainvoke({"conversation": full_conversation})

    return technical_feedback, hr_feedback
